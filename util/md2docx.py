#!/usr/bin/env python3
"""
Markdown → DOCX 报告导出脚本
用法: python md2docx.py <input.md> <output.docx>
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def convert(md_path: str, docx_path: str):
    with open(md_path, "r", encoding="utf-8") as f:
        text = f.read()

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(11)
    # 设置中文字体
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # 空行跳过
        if not line.strip():
            i += 1
            continue

        # 标题
        heading_match = re.match(r"^(#{1,6})\s+(.+)", line)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            p = doc.add_heading(title, level=level)
            i += 1
            continue

        # 图片 ![alt](path)
        img_match = re.match(r"^!\[.*\]\((.+)\)", line)
        if img_match:
            img_path = img_match.group(1)
            # 尝试解析图片路径
            resolved = _resolve_image(img_path, md_path)
            if resolved and Path(resolved).exists():
                try:
                    doc.add_picture(resolved, width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    p = doc.add_paragraph(f"[图片: {img_path}]")
                    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            else:
                p = doc.add_paragraph(f"[图片: {img_path}]")
                p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            i += 1
            continue

        # 无序列表
        ul_match = re.match(r"^[\-\*]\s+(.+)", line)
        if ul_match:
            p = doc.add_paragraph(ul_match.group(1), style="List Bullet")
            i += 1
            continue

        # 有序列表
        ol_match = re.match(r"^\d+[\.\)]\s+(.+)", line)
        if ol_match:
            p = doc.add_paragraph(ol_match.group(1), style="List Number")
            i += 1
            continue

        # 代码块
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束的 ```
            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(64, 64, 64)
            continue

        # 普通段落（支持粗体 **text** 和行内代码 `code`）
        p = doc.add_paragraph()
        _add_formatted_text(p, line)
        i += 1

    doc.save(docx_path)
    return True


def _add_formatted_text(paragraph, text):
    """解析行内的 **粗体** 和 `代码` 标记"""
    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def _resolve_image(img_path: str, md_path: str) -> str:
    """解析图片路径"""
    p = Path(img_path)
    if p.is_absolute():
        return img_path
    # 相对于 md 文件所在目录
    md_dir = Path(md_path).parent
    return str(md_dir / img_path)


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python md2docx.py <input.md> <output.docx>")
        sys.exit(1)

    convert(sys.argv[1], sys.argv[2])
    print(f"✅ 导出完成: {sys.argv[2]}")
